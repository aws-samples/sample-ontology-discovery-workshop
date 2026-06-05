# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: Apache-2.0
"""M2 — 워크샵 산출물을 파일로 내보내기.
Markdown / HTML 항상 가능. PDF는 weasyprint, docx는 python-docx(둘 다 선택적).
고객 노트북에 라이브러리가 없으면 친절히 대체 안내."""
from __future__ import annotations

import os

from .graph import OntologyGraph
from . import report as rp
from .security import audit_log, safe_export_dir, safe_export_path


def export_all(g: OntologyGraph, outdir: str = "./exports/report",
               title: str = "온톨로지 워크샵 결과",
               verified_queries: list[dict] | None = None,
               open_issues: list[str] | None = None,
               descriptions: dict | None = None,
               schema_map: dict | None = None,
               gate_data: dict | None = None,
               lang: str | None = None) -> dict:
    outdir = safe_export_dir(outdir)
    os.makedirs(outdir, exist_ok=True)
    md = rp.render_report_md(g, title, verified_queries, True, open_issues,
                             descriptions, schema_map, gate_data, lang)
    html = rp.render_report_html(g, title, verified_queries, True, open_issues,
                                 descriptions, schema_map, gate_data, lang)

    md_path = os.path.join(outdir, "workshop_report.md")
    html_path = os.path.join(outdir, "workshop_report.html")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md)
    os.chmod(md_path, 0o600)
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    os.chmod(html_path, 0o600)

    result = {"markdown": md_path, "html": html_path}

    pdf_path = os.path.join(outdir, "workshop_report.pdf")
    try:
        from weasyprint import HTML  # noqa: PLC0415
        HTML(string=html).write_pdf(pdf_path)
        os.chmod(pdf_path, 0o600)
        result["pdf"] = pdf_path
    except Exception as e:  # noqa: BLE001
        result["pdf_error"] = (
            f"PDF 생성 불가({e}). HTML을 브라우저에서 인쇄→PDF로 저장하세요.")

    docx_path = os.path.join(outdir, "workshop_report.docx")
    try:
        _export_docx(g, docx_path, title, verified_queries, open_issues,
                     descriptions, schema_map, gate_data, lang)
        os.chmod(docx_path, 0o600)
        result["docx"] = docx_path
    except Exception as e:  # noqa: BLE001
        result["docx_error"] = f"docx 생성 불가({e}). 'pip install python-docx' 필요."

    # 인계서가 참조하는 적재용 익스포트도 함께 생성
    try:
        from . import neptune_export as nx  # noqa: PLC0415
        result["neptune_cypher"] = nx.export_opencypher(
            g, os.path.join(outdir, "neptune.cypher"), schema_map)
        result["neptune_bulk"] = nx.export_bulk_csv(
            g, os.path.join(outdir, "bulk"), schema_map)
    except Exception as e:  # noqa: BLE001
        result["neptune_error"] = str(e)

    audit_log("report_files_exported", outdir=outdir, title=title, lang=lang)
    return result


def _export_docx(g: OntologyGraph, path: str, title: str,
                 verified_queries: list[dict] | None,
                 open_issues: list[str] | None = None,
                 descriptions: dict | None = None,
                 schema_map: dict | None = None,
                 gate_data: dict | None = None,
                 lang: str | None = None) -> None:
    path = safe_export_path(path)
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    md = rp.render_report_md(g, title, verified_queries, True, open_issues,
                             descriptions, schema_map, gate_data, lang)
    doc = Document()
    # 기본 폰트(한글 대응)
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10.5)

    lines = md.split("\n")
    i, n, in_fence = 0, len(lines), False
    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()

        # 코드/머메이드 펜스 내부는 통째로 스킵 (docx엔 다이어그램 렌더 불가)
        if stripped.startswith("```"):
            if stripped.startswith("```mermaid") and not in_fence:
                doc.add_paragraph("[다이어그램은 HTML/Markdown 보고서 참조]").italic = True
            in_fence = not in_fence
            i += 1
            continue
        if in_fence:
            i += 1
            continue

        # 표 블록
        if stripped.startswith("|") and i + 1 < n and rp._is_table_sep(lines[i + 1]):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            _docx_table(doc, block)
            continue

        if line.startswith("# "):
            doc.add_heading(_plain(line[2:]), level=0)
        elif line.startswith("## "):
            doc.add_heading(_plain(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(_plain(line[4:]), level=2)
        elif line.startswith("- ") or line.startswith("  - "):
            doc.add_paragraph(_plain(line.lstrip(" -")), style="List Bullet")
        elif line.startswith("> "):
            p = doc.add_paragraph(_plain(line[2:]))
            p.runs[0].italic = True
        elif stripped:
            doc.add_paragraph(_plain(line))
        i += 1
    doc.save(path)


def _docx_table(doc, rows: list[str]) -> None:
    def cells(line: str) -> list[str]:
        return [_plain(c.strip()).replace("<br/>", " ")
                for c in line.strip().strip("|").split("|")]

    header = cells(rows[0])
    body = [cells(r) for r in rows[2:]]  # rows[1] = 구분선
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        table.rows[0].cells[j].text = h
    for r in body:
        rc = table.add_row().cells
        for j in range(min(len(r), len(header))):
            rc[j].text = r[j]


def _plain(s: str) -> str:
    """docx 본문용: 마크다운 강조/코드 기호 제거."""
    import re  # noqa: PLC0415
    s = s.replace("**", "").replace("`", "")
    s = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"\1", s)  # *italic* → italic
    return s
